import re
import os
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner padding for a table cell (in twentieths of a point, dxas)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, color_hex="F2F2F2"):
    """Set background color of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_table_borders(table):
    """Set standard thin black borders for a table."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    # Configure borders
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 1/2 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC') # light gray borders
        tblBorders.append(border)
        
    tblPr.append(tblBorders)

def add_paragraph_with_runs(doc, text, style=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, space_before=0, line_spacing=1.3):
    """Add a paragraph and parse simple inline formatting like **bold** and *italic*."""
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    
    # Regex for bold and italic
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
            
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
    return p

def convert_md_to_docx(md_path, docx_path):
    print(f"Converting {md_path} to {docx_path}...")
    
    # Initialize document
    doc = docx.Document()
    
    # Page Setup (A4, 2cm margins)
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.79)  # 2cm
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)  # 3cm left margin for binding
        section.right_margin = Inches(0.79) # 2cm right
        
    # Read Markdown file
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_text = []
    
    in_table = False
    table_headers = []
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\r\n')
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
                # Write code block paragraph
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.left_indent = Inches(0.4)
                
                # We format code with Consolas font and shading
                run = p.add_run("\n".join(code_text))
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                
                # Add background formatting (xml-based) via paragraph borders or just shading
                # For simplicity, we just keep Consolas and a block indent
                code_text = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_text.append(line)
            i += 1
            continue
            
        # Handle tables
        if line.startswith('|'):
            # It's a table row
            parts = [cell.strip() for cell in line.split('|')[1:-1]]
            
            # Check if this is the separator row |---|---|
            is_separator = all(re.match(r'^:?-+:?$', part) for part in parts) if parts else False
            
            if is_separator:
                # Skip separator line
                i += 1
                continue
                
            if not in_table:
                in_table = True
                table_headers = parts
            else:
                table_rows.append(parts)
            i += 1
            continue
        else:
            if in_table:
                # We finished parsing a table, let's write it to the document
                write_table(doc, table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []
            # continue processing the line
            
        # Skip empty lines (except to end table/code block)
        if not line.strip():
            i += 1
            continue
            
        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[2:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(18)
            run.bold = True
            
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[3:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.bold = True
            
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[4:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
            
        # Lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            # Bullet list
            text = re.sub(r'^[\s\t]*[-\*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.2
            
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
                
        elif re.match(r'^\d+\.\s+', line.strip()):
            # Numbered list
            match = re.match(r'^(\d+)\.\s+(.*)', line.strip())
            num = match.group(1)
            text = match.group(2)
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.2
            
            run_num = p.add_run(f"{num}. ")
            run_num.font.name = 'Times New Roman'
            run_num.font.size = Pt(13)
            
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
                
        # Center title block on cover page
        elif line.startswith('**Khoa:**') or line.startswith('**Trường:**') or line.startswith('**Học phần:**') or line.startswith('**Giảng viên') or line.startswith('**Sinh viên') or line.startswith('**Mã số') or line.startswith('**Lớp:**') or line.startswith('**Thời gian'):
            add_paragraph_with_runs(doc, line, space_after=4)
            
        elif line.startswith('## Đề tài:') or line.startswith('### (Hybrid Search'):
            add_paragraph_with_runs(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
            
        elif line == '---':
            # Horizontal rule represents page break in our case
            doc.add_page_break()
            
        # Normal paragraphs
        else:
            align = WD_ALIGN_PARAGRAPH.LEFT
            # Justify body paragraphs
            if not line.startswith('**') and len(line) > 80:
                align = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_paragraph_with_runs(doc, line, align=align, space_after=6)
            
        i += 1
        
    if in_table:
        write_table(doc, table_headers, table_rows)
        
    # Apply global font styling to Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    # Save the document
    doc.save(docx_path)
    print("Done!")

def write_table(doc, headers, rows):
    """Generate and style a table in docx."""
    if not headers and not rows:
        return
        
    num_cols = len(headers) if headers else len(rows[0])
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Set headers
    hdr_cells = table.rows[0].cells
    for j, h_text in enumerate(headers):
        hdr_cells[j].text = ""
        p = hdr_cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(h_text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        set_cell_background(hdr_cells[j], "EAEAEA") # darker shade for header
        set_cell_margins(hdr_cells[j], top=120, bottom=120, left=150, right=150)
        hdr_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
    # Set rows
    for i, row_data in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        for j, val in enumerate(row_data):
            if j < len(row_cells):
                row_cells[j].text = ""
                p = row_cells[j].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                
                # Inline style parsing inside cells
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', val)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    else:
                        run = p.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    
                set_cell_margins(row_cells[j], top=100, bottom=100, left=150, right=150)
                row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # Alternate row shading
                if i % 2 == 1:
                    set_cell_background(row_cells[j], "FAFAFA")
                    
    # Add an empty spacing paragraph after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

if __name__ == '__main__':
    md_file = "docs/report.md"
    docx_file = "docs/report.docx"
    convert_md_to_docx(md_file, docx_file)
